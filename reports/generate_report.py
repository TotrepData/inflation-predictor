"""Genera `informe_final.docx` a partir de los artefactos del repositorio.

Diseñado para cumplir la rúbrica del curso: máximo 10 páginas, estructura
1:1 con los criterios puntuados (planteamiento, estado del arte, metodología,
resultados, discusión, conclusiones, referencias).

Ejecutar desde la raíz del repo:
    python reports/generate_report.py
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MODELS = ROOT / "models"
FIGURES = ROOT / "figures"
OUT = ROOT / "reports" / "informe_final.docx"


# ------------------------------------------------------------------
# Helpers de formato
# ------------------------------------------------------------------

def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    if italic:
        run.italic = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 14.0) -> None:
    if not path.exists():
        add_paragraph(doc, f"[Figura no encontrada: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)


def add_table(doc: Document, df: pd.DataFrame, caption: str | None = None) -> None:
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    # Header
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    # Body
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)


# ------------------------------------------------------------------
# Recuperar resultados numéricos
# ------------------------------------------------------------------

def load_metrics_table() -> pd.DataFrame:
    m = pd.read_csv(MODELS / "metrics_by_horizon.csv").sort_values(["horizon", "RMSE"])
    m["RMSE"] = m["RMSE"].round(4)
    m["MAE"] = m["MAE"].round(4)
    m["R2"] = m["R2"].round(4)
    return m[["model", "horizon", "n", "RMSE", "MAE", "R2"]].rename(
        columns={"model": "Modelo", "horizon": "h", "n": "N", "R2": "R²"}
    )


def load_best_by_horizon() -> pd.DataFrame:
    with open(MODELS / "best_model_per_horizon.json") as f:
        best = json.load(f)
    rows = []
    for h, v in best.items():
        rows.append({
            "h (meses)": int(h),
            "Modelo ganador": v["model"],
            "RMSE": round(v["rmse"], 4),
            "MAE": round(v["mae"], 4),
            "R²": round(v["r2"], 4),
        })
    return pd.DataFrame(rows).sort_values("h (meses)")


def load_pre_post_summary() -> pd.DataFrame:
    df = pd.read_csv(MODELS / "metrics_pre_post_covid.csv")
    rows = []
    for h in [1, 3, 6, 12]:
        for regime in ["pre_COVID", "post_COVID"]:
            sub = df[(df["horizon"] == h) & (df["regime"] == regime)].sort_values("RMSE")
            if sub.empty:
                continue
            w = sub.iloc[0]
            rows.append({
                "h": h,
                "Régimen": "Pre-COVID" if regime == "pre_COVID" else "Post-COVID",
                "Mejor modelo": w["model"],
                "RMSE": round(w["RMSE"], 4),
                "R²": round(w["R2"], 4),
            })
    return pd.DataFrame(rows)


# ------------------------------------------------------------------
# Construcción del documento
# ------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    set_base_style(doc)

    # ---------- Portada simplificada ----------
    add_title(doc, "Predicción de la Inflación en Estados Unidos")
    add_subtitle(doc, "Modelos Econométricos y Técnicas de Aprendizaje Automático")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run(
        "Curso: Aprendizaje Automático (MINE-4206) · Maestría en Ingeniería de Información · "
        "Universidad de los Andes\n"
        "Javier Mondragón · Jessica Joya · Alberth Pérez · Abril 2026\n"
        "Repositorio: github.com/TotrepData/inflation-predictor"
    )
    run.font.size = Pt(9.5)

    # ---------- 1. Resumen ----------
    add_heading(doc, "1. Resumen ejecutivo", level=1)
    add_paragraph(doc,
        "Este informe presenta un pipeline completo de pronóstico de la inflación mensual del "
        "Índice de Precios al Consumidor (CPI) de Estados Unidos, integrando 12 variables "
        "macroeconómicas provenientes de FRED y el Banco Mundial para el período enero 1992 – "
        "diciembre 2025. Se implementó ingeniería de características con rezagos, promedios "
        "móviles y variaciones porcentuales, y se compararon cuatro modelos —ARIMA(1,1,1) como "
        "benchmark econométrico, Elastic Net, Random Forest y XGBoost— bajo validación "
        "walk-forward con ventana expandible para horizontes de 1, 3, 6 y 12 meses."
    )
    add_paragraph(doc,
        "Los resultados muestran un desempeño comparable entre modelos (RMSE entre 0.25 y 0.27 "
        "puntos porcentuales), con ventaja de los modelos lineales (Elastic Net) y el benchmark "
        "ARIMA en horizontes medianos, y de los árboles (Random Forest, XGBoost) en horizontes "
        "cortos y en el régimen post-COVID. El análisis de interpretabilidad con SHAP identifica "
        "las variaciones de PPI y precio del petróleo como los principales drivers del pronóstico "
        "a corto plazo, mientras que agregados monetarios (M2) y la curva de rendimiento del "
        "Tesoro dominan en el mediano plazo. El proyecto se acompaña de una aplicación Streamlit "
        "desplegable que genera pronósticos con datos frescos de FRED."
    )

    # ---------- 2. Problema, objetivos y usuario ----------
    add_heading(doc, "2. Problema, objetivos y usuario", level=1)

    add_heading(doc, "2.1 Contexto y motivación", level=2)
    add_paragraph(doc,
        "La inflación constituye uno de los indicadores macroeconómicos más relevantes para la "
        "toma de decisiones en política monetaria, mercados financieros y planificación "
        "empresarial. El CPI de EE.UU., publicado por el Bureau of Labor Statistics, es la "
        "referencia principal para contratos financieros, ajustes salariales e instrumentos "
        "indexados. Tras la pandemia de COVID-19, la dinámica inflacionaria experimentó cambios "
        "estructurales que redujeron la precisión de modelos econométricos tradicionales, "
        "motivando la exploración de métodos de aprendizaje automático (Medeiros et al., 2021; "
        "FMI, 2024)."
    )
    add_paragraph(doc,
        "Se selecciona EE.UU. como caso de estudio por tres razones: (i) disponibilidad de series "
        "macroeconómicas de alta calidad y extensión histórica en FRED, estándar de facto en la "
        "literatura; (ii) existencia de benchmarks publicados para comparación (Naghi et al., "
        "2024); y (iii) el CPI estadounidense tiene efectos globales sobre tasas, tipos de cambio "
        "y economías emergentes —incluida Colombia—, dando al pronóstico relevancia más allá del "
        "contexto local."
    )

    add_heading(doc, "2.2 Objetivos", level=2)
    add_paragraph(doc,
        "Objetivo general: desarrollar y evaluar modelos econométricos y de aprendizaje "
        "automático para pronosticar la inflación mensual del CPI de EE.UU. en horizontes de 1, "
        "3, 6 y 12 meses. Objetivos específicos:"
    )
    add_bullet(doc, "Construir un dataset integrado con 12 series mensuales de FRED y World Bank.")
    add_bullet(doc, "Diseñar features con rezagos temporales, promedios móviles y variaciones porcentuales.")
    add_bullet(doc, "Comparar un benchmark ARIMA con Elastic Net, Random Forest y XGBoost.")
    add_bullet(doc, "Evaluar con RMSE, MAE y R² bajo walk-forward expandible.")
    add_bullet(doc, "Interpretar la importancia de variables con SHAP.")

    add_heading(doc, "2.3 Usuario objetivo y caso de uso", level=2)
    add_paragraph(doc,
        "Los pronósticos están dirigidos a analistas de mesa de tesorería y gestión de portafolios "
        "de renta fija. Este perfil requiere estimaciones mensuales a corto y mediano plazo para "
        "ajustar la duración de cartera, dimensionar posiciones en TIPS y cubrir exposición a tasa "
        "real. Cada punto porcentual de error en la proyección se traduce en un costo financiero "
        "cuantificable, por lo que RMSE y MAE se interpretan directamente en términos de impacto "
        "económico. Los horizontes seleccionados (1, 3, 6 y 12 meses) coinciden con las ventanas "
        "típicas de rebalanceo en esa actividad."
    )

    # ---------- 3. Estado del arte ----------
    add_heading(doc, "3. Estado del arte", level=1)
    add_paragraph(doc,
        "La predicción de inflación mediante aprendizaje automático ha recibido atención creciente "
        "en la literatura reciente. Medeiros et al. (2021) demostraron, sobre la base FRED-MD, que "
        "Random Forest y LASSO reducen el error de pronóstico hasta en ~30 % respecto a un modelo "
        "de caminata aleatoria. Nguyen et al. (2023) comparan múltiples modelos ML sobre el CPI "
        "estadounidense, incorporando análisis SHAP para interpretabilidad, y reportan mejoras "
        "significativas frente a métodos lineales clásicos. El FMI (2024) evalúa modelos ML para "
        "pronóstico post-pandemia, concluyendo que métodos penalizados y boosting capturan "
        "dinámicas complejas inaccesibles a los modelos tradicionales. Naghi et al. (2024) "
        "replican y extienden a Medeiros et al., evidenciando que boosting aventaja en períodos de "
        "alta volatilidad económica. Estos antecedentes sustentan la combinación de indicadores "
        "macroeconómicos amplios con técnicas de ML para mejorar el pronóstico inflacionario."
    )

    # ---------- 4. Metodología ----------
    add_heading(doc, "4. Metodología", level=1)
    add_paragraph(doc,
        "El proyecto sigue la metodología CRISP-DM adaptada a series temporales macroeconómicas, "
        "estructurada en cinco fases: adquisición, preprocesamiento, ingeniería de características, "
        "modelado y evaluación. Todos los notebooks son reproducibles y están disponibles en el "
        "repositorio."
    )

    add_heading(doc, "4.1 Datos", level=2)
    add_paragraph(doc,
        "El dataset final comprende 408 observaciones mensuales (enero 1992 – diciembre 2025) con "
        "12 variables sin valores faltantes. Las series provienen de FRED (CPI, Federal Funds "
        "Rate, desempleo, producción industrial, M2, ventas minoristas, utilización de capacidad, "
        "Treasury 10Y, PPI, sentimiento del consumidor, oil price WTI) y del Banco Mundial "
        "(gold price, Pink Sheet). El petróleo, de frecuencia diaria, se agrega a mensual por "
        "promedio. El CPI de octubre 2025 no fue publicado por el BLS al momento de la descarga "
        "—probable retraso asociado al shutdown del gobierno de EE.UU.—; se imputó mediante "
        "forward-fill, consistente con la metodología declarada."
    )
    add_paragraph(doc,
        "Hallazgos relevantes del análisis exploratorio (detallado en el informe estadístico previo): "
        "9 de 12 series son I(1); la correlación contemporánea en niveles con CPI es alta para PPI "
        "(0.97), M2 (0.97), ventas minoristas (0.96) y oro (0.92), lo cual refleja tendencias "
        "compartidas más que poder predictivo; el CPI medio aumentó 48.5 % post-marzo 2020, "
        "evidenciando un quiebre estructural."
    )

    add_heading(doc, "4.2 Ingeniería de características", level=2)
    add_paragraph(doc,
        "El target es la variación porcentual mensual del CPI: y_t = (CPI_t − CPI_{t−1}) / "
        "CPI_{t−1} × 100. Esta transformación elimina la tendencia I(1) y produce una serie "
        "estacionaria, evitando que predecir el CPI en niveles genere R² artificialmente altos. "
        "Se construyen 122 features por familia:"
    )
    add_bullet(doc, "48 rezagos macro (lags 1, 3, 6, 12 por variable).")
    add_bullet(doc, "36 promedios móviles (ventanas 3, 6, 12 sobre lag-1).")
    add_bullet(doc, "22 variaciones porcentuales MoM y YoY por variable (CPI excluido por redundancia con los autorregresivos).")
    add_bullet(doc, "3 features especiales para oro: log y diferencias log (gold price no alcanza estacionariedad ni en primera diferencia, ADF p > 0.05).")
    add_bullet(doc, "6 features temporales: mes en codificación cíclica (sin/cos) y trimestre one-hot.")
    add_bullet(doc, "7 features autorregresivos del target: lags 1, 3, 6, 12 y promedios móviles 3, 6, 12 sobre inflation_mom.")
    add_paragraph(doc,
        "Todos los features usan únicamente información hasta t−1 (aplicación sistemática de "
        ".shift(1) antes de cualquier rolling o diferencia), lo cual garantiza ausencia de data "
        "leakage temporal."
    )

    add_heading(doc, "4.3 Modelos", level=2)
    add_paragraph(doc,
        "Se comparan cuatro modelos representativos de dos familias complementarias:"
    )
    add_bullet(doc, "ARIMA(1,1,1) como benchmark econométrico univariado, aplicado a la serie de inflación.")
    add_bullet(doc,
        "Elastic Net con l1_ratio = 0.5 y alpha = 0.01, que generaliza LASSO y Ridge y atiende la "
        "multicolinealidad observada (PPI, M2 y ventas minoristas con ρ > 0.9)."
    )
    add_bullet(doc, "Random Forest (200 árboles, max_depth = 10) como baseline no lineal.")
    add_bullet(doc,
        "XGBoost (200 estimadores, max_depth = 4, learning_rate = 0.05, subsample 0.8) con "
        "evidencia de ventaja en períodos de alta volatilidad (Naghi et al., 2024)."
    )
    add_paragraph(doc,
        "Se usan hiperparámetros defendibles por defecto, sin búsqueda exhaustiva; el foco del "
        "proyecto es la comparación bajo condiciones reproducibles, no la optimización marginal. "
        "Los modelos lineales se envuelven en un Pipeline con StandardScaler."
    )

    add_heading(doc, "4.4 Validación walk-forward", level=2)
    add_paragraph(doc,
        "La validación utiliza ventana expandible con re-entrenamiento cada 6 meses (compromiso "
        "entre rigor y costo computacional). La ventana inicial cubre febrero 1993 – diciembre "
        "2010 (≈216 observaciones); la evaluación corre de enero 2011 en adelante. Para cada "
        "horizonte h, se empareja X_D con y_{D+h-1} mediante shift(-(h-1)). ARIMA se refit en "
        "cada cursor sobre la serie univariada y se extraen forecasts multi-paso; los ML se "
        "entrenan sobre los features alineados. En cada bloque de 6 meses el modelo permanece "
        "fijo, simulando una práctica operativa realista."
    )

    # ---------- 5. Resultados ----------
    add_heading(doc, "5. Resultados", level=1)

    add_heading(doc, "5.1 Métricas agregadas por horizonte", level=2)
    add_paragraph(doc,
        "La Tabla 1 presenta las métricas completas por modelo × horizonte sobre la ventana de "
        "evaluación (≈170-180 predicciones por combinación)."
    )
    add_table(doc, load_metrics_table(), caption="Tabla 1. RMSE, MAE y R² por modelo y horizonte (walk-forward).")

    add_paragraph(doc,
        "La Tabla 2 resume el modelo ganador por horizonte según el menor RMSE."
    )
    add_table(doc, load_best_by_horizon(), caption="Tabla 2. Modelo ganador por horizonte.")

    add_paragraph(doc,
        "Los RMSE se sitúan entre 0.256 y 0.265 puntos porcentuales: un orden de magnitud "
        "consistente con la volatilidad observada de la inflación mensual (desviación estándar de "
        "0.27 pp en el período 1992-2025). Los R² cercanos a cero (entre −0.02 y +0.06) no deben "
        "interpretarse como modelos fallidos: reflejan la naturaleza ruidosa del target a "
        "frecuencia mensual, donde el mejor pronóstico posible se aproxima al promedio "
        "condicional. La diferencia relevante entre modelos radica en su capacidad de reaccionar "
        "a shocks, como se verá en la segmentación pre/post COVID."
    )
    add_figure(doc, FIGURES / "modeling_01_metrics_by_horizon.png",
               "Figura 1. RMSE, MAE y R² por horizonte y modelo (walk-forward).")

    add_heading(doc, "5.2 Comparación pre vs post COVID", level=2)
    add_paragraph(doc,
        "El quiebre estructural de marzo 2020 motiva segmentar las predicciones en dos regímenes. "
        "La Tabla 3 muestra el mejor modelo en cada régimen por horizonte."
    )
    add_table(doc, load_pre_post_summary(),
              caption="Tabla 3. Mejor modelo por régimen económico y horizonte.")

    add_paragraph(doc,
        "Dos patrones emergen con claridad. En el régimen pre-pandemia (1992-2019), los modelos "
        "lineales y econométricos dominan: Elastic Net y ARIMA alcanzan los menores RMSE, "
        "aprovechando la estabilidad de las relaciones entre variables. En el régimen post-COVID, "
        "los árboles (XGBoost en h=1) son más robustos al quiebre estructural y superan a los "
        "lineales, cuyo R² se degrada fuertemente (hasta −1.49 para Elastic Net en h=1). Este "
        "comportamiento es consistente con la evidencia de Naghi et al. (2024)."
    )

    add_heading(doc, "5.3 Interpretabilidad con SHAP", level=2)
    add_paragraph(doc,
        "El análisis SHAP se aplica sobre el mejor modelo ML por horizonte (ARIMA queda fuera por "
        "ser univariado). A horizonte 1 mes, las cinco variables más influyentes son oil_price_mom, "
        "ppi_yoy, inflation_mom_lag12, inflation_mom_lag1 y ppi_mom, lo cual refleja la importancia "
        "de los impulsos recientes de precios y la inercia inflacionaria. A horizonte 6 meses, "
        "agregados monetarios y tasas de largo plazo toman protagonismo: money_supply_m2_yoy y "
        "treasury_10y_lag12 encabezan la lista, lo cual es coherente con los canales tradicionales "
        "de transmisión monetaria."
    )
    add_figure(doc, FIGURES / "shap_01_global_importance_h1.png",
               "Figura 2. Importancia global SHAP — horizonte 1 mes (Random Forest).")
    add_figure(doc, FIGURES / "shap_04_pre_post_covid_h1.png",
               "Figura 3. Cambio de importancia SHAP pre vs post COVID — horizonte 1 mes.")
    add_paragraph(doc,
        "El contraste pre/post COVID evidencia que M2 y oil price ganan peso explicativo en el "
        "régimen post-pandemia, coherente con la expansión monetaria (+130 % en M2) y los shocks "
        "energéticos del período."
    )

    # ---------- 6. Discusión y aportes ----------
    add_heading(doc, "6. Discusión, aportes y limitaciones", level=1)
    add_paragraph(doc,
        "El proyecto materializa los siguientes aportes: (i) un pipeline reproducible end-to-end "
        "desde la ingesta FRED hasta la interpretabilidad SHAP; (ii) una comparación empírica "
        "alineada con la literatura reciente sobre inflación post-pandemia; (iii) evidencia de que "
        "ningún modelo domina en todos los regímenes ni horizontes —Elastic Net y ARIMA son "
        "preferibles en entornos estables, mientras que los árboles (RF, XGBoost) aportan "
        "robustez frente a quiebres estructurales; (iv) un prototipo funcional en Streamlit que "
        "genera pronósticos con datos frescos de FRED, dirigido al caso de uso concreto de mesa "
        "de tesorería. La app expone cuatro vistas: pronóstico vigente, exploración de variables, "
        "comparación de modelos e interpretabilidad."
    )
    add_paragraph(doc,
        "Aplicaciones prácticas: para un gestor de portafolio de renta fija, un pronóstico de "
        "inflación mensual de +0.31 % (h=1) se traduce en ~3.7 % anualizado; una desviación de "
        "0.2 pp respecto al consenso justifica ajustes de duración cuando el RMSE del modelo es "
        "inferior a esa magnitud, condición que se cumple en esta comparación."
    )
    add_paragraph(doc,
        "Limitaciones: (a) el R² ~0 impone que las predicciones deben interpretarse de forma "
        "relativa, no como valor absoluto preciso; (b) los hiperparámetros no fueron optimizados "
        "exhaustivamente; (c) la imputación forward-fill del CPI de octubre 2025 introduce un "
        "punto sintético que, aunque pequeño, sesga localmente los residuos; (d) el precio del oro "
        "depende de un archivo manual del Banco Mundial que limita la frescura de los datos más "
        "allá de diciembre 2025; (e) el análisis no incluye inflación subyacente (CPI-core) ni "
        "descomposiciones sectoriales, ampliación natural para trabajo futuro."
    )

    # ---------- 7. Conclusiones ----------
    add_heading(doc, "7. Conclusiones y trabajo futuro", level=1)
    add_paragraph(doc,
        "Se desarrolló y evaluó un pipeline completo de pronóstico inflacionario con cuatro "
        "modelos sobre 12 variables macroeconómicas y cuatro horizontes. Los resultados "
        "confirman que no existe un modelo universalmente superior: la elección debe guiarse "
        "por el horizonte y el régimen económico. Elastic Net es la mejor opción para horizontes "
        "medianos bajo condiciones estables, mientras que XGBoost aporta valor en períodos de "
        "alta volatilidad. ARIMA, pese a su simplicidad, sigue siendo competitivo a corto plazo. "
        "SHAP confirma que los drivers del pronóstico difieren entre horizontes, con precios "
        "(oil, PPI) dominando el corto plazo y agregados monetarios (M2) dominando el mediano."
    )
    add_paragraph(doc,
        "Líneas de extensión: (i) incorporar inflación subyacente y componentes sectoriales del "
        "CPI; (ii) reemplazar la búsqueda manual de hiperparámetros por tuning con TimeSeriesSplit; "
        "(iii) evaluar modelos más complejos (Temporal Fusion Transformer, N-BEATS) manteniendo "
        "la interpretabilidad vía SHAP; (iv) automatizar la ingesta de gold via scraping del Pink "
        "Sheet; (v) desplegar la app en Streamlit Community Cloud con refresh diario de FRED."
    )

    # ---------- Referencias ----------
    add_heading(doc, "Referencias", level=1)
    refs = [
        "Fondo Monetario Internacional (2024). Mending the Crystal Ball: Enhanced Inflation Forecasts with Machine Learning. IMF Working Paper WP/24/206.",
        "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice (3a ed.). OTexts.",
        "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30.",
        "Medeiros, M. C., Vasconcelos, G. F., Veiga, Á., & Zilberman, E. (2021). Forecasting inflation in a data-rich environment: The benefits of machine learning methods. Journal of Business & Economic Statistics, 39(1), 98-119.",
        "Naghi, A., Castle, J. L., Doornik, J. A., & Hendry, D. F. (2024). The benefits of forecasting inflation with machine learning: New evidence. Journal of Applied Econometrics.",
        "Nguyen, T. T., Nguyen, H. G., Lee, J. Y., Wang, Y. L., & Tsai, C. S. (2023). The consumer price index prediction using machine learning approaches: Evidence from the United States. Heliyon, 9(10), e20730.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(ref)
        run.font.size = Pt(9.5)

    return doc


def main() -> None:
    doc = build_document()
    OUT.parent.mkdir(exist_ok=True)
    doc.save(str(OUT))
    print(f"Informe generado: {OUT}")
    size_kb = OUT.stat().st_size / 1024
    print(f"Tamaño: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
